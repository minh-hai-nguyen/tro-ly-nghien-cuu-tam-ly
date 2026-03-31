# -*- coding: utf-8 -*-
"""Tạo tài liệu hướng dẫn sử dụng web Trợ lý Nghiên cứu"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2)

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)

def p(text, bold=False, italic=False, size=12):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return para

def bullet(text, size=12):
    para = doc.add_paragraph(text, style='List Bullet')
    for r in para.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(size)

def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

# ===== BÌA =====
doc.add_paragraph()
title = doc.add_heading('', level=0)
r = title.add_run('TRỢ LÝ NGHIÊN CỨU TÂM LÝ HỌC\nHướng dẫn sử dụng')
r.font.name = 'Times New Roman'
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Chủ đề: Lo Âu ở Học Sinh THCS và THPT\nViệt Nam \u2014 Đông Nam Á \u2014 Thế giới\n\nPhiên bản 1.0 \u2014 Tháng 3/2026')
r.font.name = 'Times New Roman'
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ===== MỤC LỤC =====
h('Mục lục')
p('1. Giới thiệu tổng quan')
p('2. Cách truy cập')
p('3. Tab Hỏi đáp \u2014 Chatbot RAG')
p('4. Tab Tổng hợp liên bài')
p('5. Tab Thư viện')
p('6. Cơ sở dữ liệu bài báo')
p('7. Cách đọc kết quả tìm kiếm')
p('8. Câu hỏi gợi ý')
p('9. Hạn chế và lưu ý')
p('10. Hướng phát triển tương lai')
doc.add_page_break()

# ===== 1. GIỚI THIỆU =====
h('1. Giới thiệu tổng quan')
p('Trợ lý Nghiên cứu Tâm lý học là công cụ web hỗ trợ nghiên cứu về sức khỏe tâm thần ở học sinh. Công cụ sử dụng công nghệ RAG (Retrieval-Augmented Generation \u2014 Tìm kiếm Tăng cường) để tìm kiếm thông tin từ cơ sở dữ liệu gồm 10 bài báo nghiên cứu đã được dịch sang tiếng Việt.')

p('Đặc điểm chính:', bold=True)
bullet('Tìm kiếm thông minh bằng tiếng Việt và tiếng Anh')
bullet('Trả lời có trích dẫn nguồn cụ thể (bài báo, trang)')
bullet('Bấm vào nguồn để xem ngay bản dịch hoặc bản gốc PDF')
bullet('Tổng hợp liên bài \u2014 so sánh kết quả giữa nhiều nghiên cứu')
bullet('Thư viện tài liệu \u2014 tải PDF và DOCX')
bullet('Không cần đăng nhập \u2014 truy cập tự do')

p('Phạm vi nghiên cứu:', bold=True)
bullet('Việt Nam: Hà Nội, Lạng Sơn (DTTS), Huế (cơ sở bảo trợ), quốc gia (V-NAMHS)')
bullet('Đông Nam Á: 10 nước ASEAN (GBD 2021), Philippines')
bullet('Châu Á: Trung Quốc (Suzhou), Nam Á (8 nước)')
bullet('Thế giới: Tổng quan tường thuật 61 bài, can thiệp tại LMIC')

doc.add_page_break()

# ===== 2. CÁCH TRUY CẬP =====
h('2. Cách truy cập')
p('Bước 1: Mở trình duyệt (Chrome, Firefox, Edge)')
p('Bước 2: Nhập địa chỉ http://localhost:8000')
p('Bước 3: Trang web mở ra với 3 tab: Hỏi đáp, Tổng hợp liên bài, Thư viện')
p('Lưu ý: Server phải đang chạy trên máy. Khởi động server bằng lệnh:', italic=True)
p('cd Lo-au/web && python -c "import uvicorn; from app import app; uvicorn.run(app, host=\'0.0.0.0\', port=8000)"', size=10)

doc.add_page_break()

# ===== 3. TAB HỎI ĐÁP =====
h('3. Tab Hỏi đáp \u2014 Chatbot RAG')
p('Đây là chức năng chính. Nhập câu hỏi bằng tiếng Việt hoặc tiếng Anh, hệ thống sẽ tìm kiếm trong 1.410 đoạn văn từ 40 tài liệu và trả về các đoạn liên quan nhất.')

p('Mỗi kết quả bao gồm:', bold=True)
bullet('Tiêu đề bài báo (bấm để xem PDF)')
bullet('Phần trăm liên quan (ví dụ: 82%)')
bullet('Số trang tham chiếu trong bài gốc (ví dụ: Trang 5)')
bullet('Đoạn trích dẫn liên quan')
bullet('Các nút: Xem bản dịch | Tải DOCX | Xem bản gốc | Link gốc')

p('Cách đọc kết quả:', bold=True)
bullet('Độ liên quan > 70%: Rất phù hợp \u2014 đoạn trích trả lời trực tiếp câu hỏi')
bullet('Độ liên quan 50-70%: Phù hợp một phần \u2014 đoạn trích có liên quan nhưng không trả lời trực tiếp')
bullet('Độ liên quan < 50%: Ít liên quan \u2014 cần đặt câu hỏi khác')

p('Mẹo đặt câu hỏi tốt:', bold=True)
bullet('Cụ thể: "Tỷ lệ lo âu ở HS THPT Hà Nội" thay vì "lo âu"')
bullet('Đề cập công cụ: "GAD-7 là gì?" hoặc "DASS-21 cho kết quả gì?"')
bullet('So sánh: "So sánh tỷ lệ lo âu giữa Việt Nam và Philippines"')
bullet('Yếu tố: "Thiếu ngủ ảnh hưởng đến trầm cảm như thế nào?"')

doc.add_page_break()

# ===== 4. TAB TỔNG HỢP =====
h('4. Tab Tổng hợp liên bài')
p('Tab này trình bày 6 bảng tổng hợp kết quả từ nhiều bài báo \u2014 cho phép so sánh liên quốc gia, liên thời gian, liên phương pháp.')

p('6 bảng tổng hợp:', bold=True)

t = doc.add_table(rows=7, cols=3)
t.style = 'Table Grid'
headers = ['STT', 'Bảng tổng hợp', 'Số bài gộp']
for i, hdr in enumerate(headers):
    c = t.rows[0].cells[i]
    c.text = hdr
    for p2 in c.paragraphs:
        for r in p2.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
    shade(c, 'D9E2F3')

data = [
    ('1', 'Tỷ lệ lo âu/trầm cảm theo quốc gia và năm', '8 bài'),
    ('2', 'Sàng lọc vs Chẩn đoán \u2014 Khoảng cách quan trọng', '5+ bài'),
    ('3', 'Yếu tố nguy cơ và bảo vệ (cross-study)', '7 bài'),
    ('4', 'Xu hướng theo thời gian (1990\u20132025)', '4 bài'),
    ('5', 'Hiệu quả can thiệp tại trường học', '2 bài'),
    ('6', 'Khoảng trống dịch vụ SKTT', '3 bài'),
]
for ri, (a, b, c) in enumerate(data):
    for ci, v in enumerate([a, b, c]):
        cell = t.rows[ri+1].cells[ci]
        cell.text = v
        for p2 in cell.paragraphs:
            for r in p2.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)

doc.add_paragraph()
p('Mỗi ô trong bảng có link bấm trực tiếp \u2014 mở PDF bài báo tương ứng ở khung bên phải.', italic=True)

doc.add_page_break()

# ===== 5. THƯ VIỆN =====
h('5. Tab Thư viện')
p('Liệt kê tất cả tài liệu có sẵn (28 file PDF + DOCX). Bấm vào để xem PDF trực tiếp hoặc tải DOCX.')
p('Các loại tài liệu:', bold=True)
bullet('Bản dịch tiếng Việt (DICH_*.pdf) \u2014 đầy đủ từ bài gốc')
bullet('Bản gốc tiếng Anh (GOC_*.pdf) \u2014 PDF từ tạp chí')
bullet('Báo cáo tổng hợp (Lo au - *.pdf) \u2014 do nhóm NC viết')
bullet('Bản dịch DOCX (DICH_*.docx) \u2014 chỉnh sửa được')

doc.add_page_break()

# ===== 6. CSDL =====
h('6. Cơ sở dữ liệu bài báo')
p('Hệ thống hiện có 10 bài nghiên cứu chính đã dịch đầy đủ:', bold=True)

t2 = doc.add_table(rows=11, cols=4)
t2.style = 'Table Grid'
for i, hdr in enumerate(['#', 'Tác giả (Năm)', 'Nội dung', 'Phạm vi']):
    c = t2.rows[0].cells[i]
    c.text = hdr
    for p2 in c.paragraphs:
        for r in p2.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
    shade(c, 'D9E2F3')

papers = [
    ('01', 'Jenkins (2023)', 'Trầm cảm lo âu HS THCS đa sắc tộc', 'Mỹ'),
    ('02', 'Hoa (2024)', 'Lo âu HS THPT Hà Nội sau COVID', 'VN Hà Nội'),
    ('03', 'GBD ASEAN (2025)', 'Dịch tễ 10 RLTT tại ASEAN', '10 nước ĐNA'),
    ('04', 'Zhameden (2025)', 'Can thiệp trường học tại LMIC', 'LMIC'),
    ('05', 'Anderson (2025)', 'Yếu tố gia tăng lo âu VTN', 'Thế giới'),
    ('06', 'V-NAMHS (2022)', 'Khảo sát SKTT quốc gia VN', 'VN 38 tỉnh'),
    ('07', 'Zhu (2025)', 'Trầm cảm HS THCS/THPT Trung Quốc', 'TQ Suzhou'),
    ('08', 'Mudunna (2025)', 'SKTT VTN Nam Á, tổng quan', 'Nam Á 8 nước'),
    ('09', 'Puyat (2025)', 'Trầm cảm thanh niên Philippines', 'Philippines'),
    ('10', 'Pham (2024)', 'SKTT VTN cơ sở bảo trợ xã hội', 'VN Huế'),
]
for ri, (a, b, c, d) in enumerate(papers):
    for ci, v in enumerate([a, b, c, d]):
        cell = t2.rows[ri+1].cells[ci]
        cell.text = v
        for p2 in cell.paragraphs:
            for r in p2.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(9)

doc.add_paragraph()
p('Ngoài ra còn 30+ bài PDF gốc khác trong thư viện, chưa dịch nhưng đã được index vào RAG.', italic=True)

doc.add_page_break()

# ===== 7. CÁCH ĐỌC KẾT QUẢ =====
h('7. Cách đọc kết quả tìm kiếm')
p('Khi nhận kết quả từ chatbot, chú ý:', bold=True)
bullet('Badge xanh "82%" = độ liên quan giữa câu hỏi và đoạn trích')
bullet('Badge cam "Trang 5" = vị trí chính xác trong bài gốc')
bullet('Bấm tiêu đề xanh = mở PDF bản dịch ở khung bên phải')
bullet('"Xem bản dịch" = PDF tiếng Việt đầy đủ')
bullet('"Xem bản gốc" = PDF tiếng Anh từ tạp chí')
bullet('"Link gốc" = trang web tạp chí (DOI)')

p('Lưu ý quan trọng:', bold=True)
bullet('Hệ thống TÌM KIẾM chứ không TẠO NỘI DUNG \u2014 mọi thông tin đều trích từ bài báo thực')
bullet('Luôn kiểm tra nguồn gốc trước khi trích dẫn trong nghiên cứu')
bullet('Số liệu sàng lọc (DASS-21, GAD-7) và chẩn đoán (DISC-5) rất khác nhau \u2014 xem bảng 2 trong tab Tổng hợp')

doc.add_page_break()

# ===== 8. CÂU HỎI GỢI Ý =====
h('8. Câu hỏi gợi ý')
p('Dịch tễ học:', bold=True)
bullet('"Tỷ lệ lo âu ở học sinh Việt Nam là bao nhiêu?"')
bullet('"So sánh tỷ lệ trầm cảm giữa Việt Nam, Philippines và Trung Quốc"')
bullet('"Nhóm 10-14 tuổi ở ASEAN có gánh nặng bệnh tật như thế nào?"')
bullet('"Tỷ lệ lo âu ở học sinh dân tộc thiểu số Lạng Sơn?"')

p('Yếu tố nguy cơ và bảo vệ:', bold=True)
bullet('"Thiếu ngủ ảnh hưởng đến trầm cảm như thế nào?"')
bullet('"Chăm sóc cảm xúc có quan trọng hơn chăm sóc thể chất không?"')
bullet('"Yếu tố nào bảo vệ học sinh khỏi lo âu?"')
bullet('"LGBTQ+ có tỷ lệ trầm cảm cao hơn không?"')

p('Phương pháp nghiên cứu:', bold=True)
bullet('"GAD-7 là gì? So sánh với DASS-21"')
bullet('"DISC-5 khác gì với thang sàng lọc?"')
bullet('"Sự khác biệt giữa sàng lọc và chẩn đoán lo âu?"')

p('Can thiệp:', bold=True)
bullet('"Can thiệp CBT tại trường học có hiệu quả không?"')
bullet('"Tại sao can thiệp hiệu quả cho trầm cảm nhưng kém cho lo âu?"')
bullet('"Có RCT nào về SKTT tại trường học Việt Nam không?"')

p('Dịch vụ SKTT:', bold=True)
bullet('"Bao nhiêu phần trăm VTN Việt Nam tiếp cận dịch vụ SKTT?"')
bullet('"Phụ huynh Việt Nam có nhận ra con cần giúp đỡ không?"')

doc.add_page_break()

# ===== 9. HẠN CHẾ =====
h('9. Hạn chế và lưu ý')
bullet('Hệ thống chỉ tìm kiếm trong 40 tài liệu đã có \u2014 không trả lời câu hỏi ngoài phạm vi')
bullet('Kết quả là TÌM KIẾM (retrieval), không phải TRẢ LỜI (generation) \u2014 đoạn trích có thể không trả lời trực tiếp câu hỏi')
bullet('Độ liên quan phụ thuộc vào cách đặt câu hỏi \u2014 câu hỏi quá chung sẽ cho kết quả kém')
bullet('Một số bài chỉ có bản gốc tiếng Anh (chưa dịch)')
bullet('Bảng tổng hợp liên bài được xây dựng thủ công \u2014 cần cập nhật khi thêm bài mới')
bullet('Server cần chạy trên máy cục bộ (localhost) \u2014 chưa deploy online')

doc.add_page_break()

# ===== 10. HƯỚNG PHÁT TRIỂN =====
h('10. Hướng phát triển tương lai')

p('A. Nâng cao chất lượng RAG:', bold=True)
bullet('Tích hợp LLM (Large Language Model) để tổng hợp câu trả lời thay vì chỉ trích dẫn đoạn văn')
bullet('Hỏi đáp đa lượt (multi-turn) \u2014 chatbot nhớ ngữ cảnh cuộc hội thoại')
bullet('Tự động cập nhật RAG khi thêm bài mới vào thư mục')

p('B. Mở rộng cơ sở dữ liệu:', bold=True)
bullet('Thêm bài từ tạp chí Việt Nam (Tạp chí Y học, Tạp chí Tâm lý học)')
bullet('Tích hợp API PubMed/Google Scholar để tìm bài mới tự động')
bullet('Hỗ trợ upload bài mới qua giao diện web')

p('C. Phân tích nâng cao:', bold=True)
bullet('Tự động tạo bảng tổng hợp liên bài khi thêm bài mới')
bullet('Forest plot tương tác \u2014 so sánh tỷ lệ giữa các NC trên biểu đồ')
bullet('Network analysis \u2014 bản đồ quan hệ giữa các yếu tố nguy cơ/bảo vệ')
bullet('Timeline \u2014 xu hướng SKTT VTN theo thời gian trên biểu đồ tương tác')

p('D. Hỗ trợ viết nghiên cứu:', bold=True)
bullet('Tự động tạo bảng "Tóm tắt các nghiên cứu" cho báo cáo tổng hợp')
bullet('Đề xuất khoảng trống nghiên cứu (gap) dựa trên cơ sở dữ liệu')
bullet('Hỗ trợ viết phần Literature Review theo phong cách hàn lâm')
bullet('Tạo đề cương nghiên cứu dựa trên gap đã xác định')

p('E. Deploy và chia sẻ:', bold=True)
bullet('Deploy lên Render/Vercel để truy cập từ mọi nơi')
bullet('Mua tên miền riêng (ví dụ: loauhocsinh.vn)')
bullet('Cho phép nhiều nhóm nghiên cứu cùng sử dụng')
bullet('Hỗ trợ đa chủ đề (không chỉ lo âu \u2014 mở rộng sang trầm cảm, stress, bắt nạt, v.v.)')

p('F. Tự động hóa:', bold=True)
bullet('Cron job tự động quét bài mới trên PubMed hàng tuần')
bullet('Tự động dịch bài mới bằng AI + kiểm tra chất lượng')
bullet('Alert email khi có bài mới liên quan đến chủ đề nghiên cứu')
bullet('Tích hợp Zotero/Mendeley để đồng bộ thư viện tham khảo')

doc.add_paragraph()
p('Tài liệu này được tạo tự động ngày 29/03/2026.', italic=True, size=10)

doc.save('Huong dan su dung - Tro ly Nghien cuu.docx')
