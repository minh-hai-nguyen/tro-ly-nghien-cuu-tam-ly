# -*- coding: utf-8 -*-
"""
Rebuild RAG database hoàn toàn — 35 bài + cross-study + insights
Idempotent — chạy lại bao nhiêu lần cũng được.
"""
import os, sys, json
os.environ['PYTHONIOENCODING'] = 'utf-8'
import chromadb
import docx

BASE = os.path.dirname(os.path.abspath(__file__))
DICH_DIR = os.path.join(BASE, '..', '03_Ban-dich')
TT_DIR = os.path.join(BASE, '..', 'Tom-tat-tung-bai')
RAG_DIR = os.path.join(BASE, 'rag_db_full')

def read_docx(fp, max_chars=4000):
    try:
        doc = docx.Document(fp)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        for t in doc.tables:
            for row in t.rows:
                text += '\n' + ' | '.join([c.text.strip()[:50] for c in row.cells])
        return text[:max_chars]
    except:
        return ''

# ===== PAPER DEFINITIONS =====
papers = [
    # VN
    {'id':'VN01','region':'Việt Nam','region_code':'VN','file_dich':'02_Hoa_2024_Frontiers.docx','file_tt':'VN1_Hoa_2024.docx',
     'title':'Lo âu HS THPT Hà Nội sau COVID-19','authors':'Hoa et al. 2024','journal':'Frontiers in Public Health Q1',
     'n':'3910','location':'Hà Nội','tool':'GAD-7','key':'Lo âu 40,6%; nữ>nam; phương pháp hỗn hợp; GAD-7 α=0,916'},
    {'id':'VN02','region':'Việt Nam','region_code':'VN','file_dich':'06_VNAMHS_2022.docx','file_tt':'VN2_VNAMHS_2022.docx',
     'title':'V-NAMHS — Khảo sát quốc gia SKTT VTN VN','authors':'V-NAMHS 2022','journal':'Báo cáo quốc gia',
     'n':'5996','location':'38 tỉnh VN','tool':'DISC-5/DSM-5','key':'Chẩn đoán 2,3%; tiếp cận 8,4%; phụ huynh nhận ra 5,1%'},
    {'id':'VN03','region':'Việt Nam','region_code':'VN','file_dich':'10_Pham_2024_VN_SocialSupport.docx','file_tt':'VN3_Pham_2024.docx',
     'title':'Hỗ trợ xã hội và SKTT VTN SSF Huế','authors':'Pham et al. 2024','journal':'J Affective Disorders Q1',
     'n':'500','location':'Huế','tool':'DASS-21','key':'Hỗ trợ xã hội bảo vệ lo âu'},
    {'id':'VN14','region':'Việt Nam','region_code':'VN','file_dich':'14_HoangTrungHoc_2025_AJPR.docx','file_tt':'VN14_HoangTrungHoc_2025.docx',
     'title':'SKTT VTN VN trước và sau COVID-19','authors':'Hoàng Trung Học 2025','journal':'AJPR',
     'n':'2000','location':'VN toàn quốc','tool':'DASS-21','key':'Lo âu 41,5%→25,4% phục hồi; game điện tử Beta=0,176'},
    {'id':'VN15','region':'Việt Nam','region_code':'VN','file_dich':'15_NgoAnhVinh_2024_JAffectDisordReports.docx','file_tt':'VN15_NgoAnhVinh_2024.docx',
     'title':'SKTT VTN DTTS Lạng Sơn','authors':'Ngô Anh Vinh et al. 2024','journal':'J Affect Disord Reports',
     'n':'845','location':'Lạng Sơn DTTS','tool':'DASS-21+ACEs','key':'Lo âu 54,4%; ACEs liên quan; DTTS dễ bị tổn thương'},
    {'id':'VN16','region':'Việt Nam','region_code':'VN','file_dich':'16_BaoQuyen_2025_YHCD.docx','file_tt':'VN16_BaoQuyen_2025.docx',
     'title':'SKTT HS THPT Hà Nội — DASS-21','authors':'Bảo Quyên et al. 2025','journal':'TC Y học Cộng đồng',
     'n':'501','location':'Hà Nội','tool':'DASS-21','key':'Lo âu 86,2% CAO NHẤT; 78% nữ; VinUni'},
    {'id':'VN17','region':'Việt Nam','region_code':'VN','file_dich':'17_NguyenDanhLam_2022_YHVN.docx','file_tt':'VN17_DanhLam_2022.docx',
     'title':'Stress, lo âu, trầm cảm HS THPT Yên Định Thanh Hóa','authors':'Nguyễn Danh Lâm 2022','journal':'TC Y học VN',
     'n':'482','location':'Yên Định, Thanh Hóa','tool':'DASS-21','key':'Lo âu 49,0%; tự hại 10%; tự tử 1,4%'},
    {'id':'VN18','region':'Việt Nam','region_code':'VN','file_dich':'18_AnGiang_2025_YHVN.docx','file_tt':'VN18_AnGiang_2025.docx',
     'title':'Sàng lọc DASS-21 HS Long Bình An Giang','authors':'An Giang 2025','journal':'TC Y học VN',
     'n':'366','location':'Long Bình, An Giang','tool':'DASS-21','key':'Lo âu 61,2%; trầm cảm 47,3%; stress 38,0%'},
    {'id':'VN19','region':'Việt Nam','region_code':'VN','file_dich':'19_TranThaoVi_2025_TLH.docx','file_tt':'VN19_TranThaoVi_2025.docx',
     'title':'Lo âu, trầm cảm và lạc quan ở VTN Huế','authors':'Hồ Thị Trúc Quỳnh 2025','journal':'TC Tâm lý học',
     'n':'685','location':'Huế','tool':'DASS-21+LOT-R','key':'Lo âu 65,8%; lạc quan trung gian β gián tiếp=−0,24'},
    {'id':'VN20','region':'Việt Nam','region_code':'VN','file_dich':'20_TranHoVinhLoc_2024_YHTPHCM.docx','file_tt':'VN20_TranHoVinhLoc_2024.docx',
     'title':'DAS và yếu tố liên quan HS THPT TPHCM','authors':'Trần Hồ Vĩnh Lộc 2024','journal':'TC Y học TPHCM',
     'n':'976','location':'TPHCM','tool':'DASS-Y','key':'Lo âu 25,1% DASS-Y thấp hơn DASS-21; áp lực ESSA; cha mẹ ly hôn'},
    # Đông Nam Á
    {'id':'QT06','region':'Đông Nam Á','region_code':'SEA','file_dich':'08_Mudunna_2025_LancetSEA.docx','file_tt':'',
     'title':'GBD RLTT ASEAN 10 nước 1990-2021','authors':'Mudunna/GBD 2025','journal':'Lancet SEA',
     'n':'GBD','location':'ASEAN 10 nước','tool':'GBD 2021','key':'80,4 triệu ca RLTT; tăng 70%; VN đứng thứ 3'},
    {'id':'QT07','region':'Đông Nam Á','region_code':'SEA','file_dich':'09_Puyat_2025_Filipino.docx','file_tt':'',
     'title':'SKTT VTN Philippines','authors':'Puyat 2025','journal':'Phil J Health',
     'n':'quốc gia','location':'Philippines','tool':'GSHS','key':'Trầm cảm 9,6%→20,9% gấp đôi; tiếp cận chỉ 2%'},
    {'id':'QT10','region':'Đông Nam Á','region_code':'SEA','file_dich':'03_GBD_ASEAN_2025_Lancet.docx','file_tt':'',
     'title':'GBD ASEAN Lancet — 10 RLTT','authors':'GBD ASEAN 2025','journal':'Lancet IF>160',
     'n':'GBD','location':'ASEAN','tool':'GBD 2021','key':'VN đứng thứ 3 ASEAN gánh nặng'},
    # Đông Bắc Á
    {'id':'QT05','region':'Đông Bắc Á','region_code':'NEA','file_dich':'07_Zhu_2025_BMC_China.docx','file_tt':'',
     'title':'Lo âu/trầm cảm HS TQ (Suzhou)','authors':'Zhu et al. 2025','journal':'BMC Public Health Q1',
     'n':'5000+','location':'Suzhou, TQ','tool':'DASS-21','key':'Giấc ngủ <5h AOR=13,71; gia đình đơn thân AOR=1,434'},
    {'id':'QT08','region':'Đông Bắc Á','region_code':'NEA','file_dich':'12_Wen_2020_IJERPH.docx','file_tt':'',
     'title':'LPA lo âu VTN nông thôn TQ','authors':'Wen et al. 2020','journal':'IJERPH Q1',
     'n':'3000+','location':'TQ nông thôn','tool':'DASS-21+LPA','key':'Áp lực OR=11,58; hỗ trợ trường OR=0,562'},
    {'id':'QT09','region':'Đông Bắc Á','region_code':'NEA','file_dich':'13_Xu_2021_JAffectDisord.docx','file_tt':'',
     'title':'Trầm cảm lo âu VTN TQ','authors':'Xu et al. 2021','journal':'J Affect Disord Q1',
     'n':'2000+','location':'TQ','tool':'DASS-21','key':'Lo âu liên quan gia đình, MXH'},
    {'id':'QT34','region':'Đông Bắc Á','region_code':'NEA','file_dich':'34_Korea_MH_Trends_2024.docx','file_tt':'QT34_Korea_Trends.docx',
     'title':'Xu hướng SKTT VTN Hàn Quốc theo thu nhập 2006-2022','authors':'Cho et al. 2024','journal':'Nat Sci Rep Q1',
     'n':'KYRBS quốc gia','location':'Hàn Quốc','tool':'KYRBS','key':'Đảo chiều COVID; nghèo 62,77% vs giàu 40,07%; chênh 22,7'},
    # Châu Á khác
    {'id':'QT02','region':'Châu Á','region_code':'ASIA','file_dich':'11_Saikia_2023_IJCM.docx','file_tt':'',
     'title':'Lo âu trầm cảm HS Ấn Độ','authors':'Saikia 2023','journal':'IJCM',
     'n':'400','location':'Assam, Ấn Độ','tool':'DASS-21','key':'Lo âu 30,5%'},
    {'id':'QT11','region':'Châu Á','region_code':'ASIA','file_dich':'','file_tt':'',
     'title':'DASS-21 Chandigarh Ấn Độ','authors':'Bhardwaj 2020','journal':'—',
     'n':'~300','location':'Chandigarh, Ấn Độ','tool':'DASS-21','key':'Lo âu/trầm cảm HS Ấn Độ'},
    # Châu Âu
    {'id':'QT21','region':'Âu-Mỹ-Úc','region_code':'WEST','file_dich':'21_Norway_2025_SocSciMed.docx','file_tt':'QT21_Norway_2025.docx',
     'title':'Giải thích xu hướng tăng distress VTN Na Uy 2011-2024','authors':'Brunborg et al. 2025','journal':'SocSciMed Q1',
     'n':'979043','location':'Na Uy','tool':'HSCL-6','key':'Bất mãn trường=giải thích chính; MXH một phần; decomposition'},
    {'id':'QT24','region':'Âu-Mỹ-Úc','region_code':'WEST','file_dich':'24_WHO_Europe_2025_LancetRegional.docx','file_tt':'QT24_WHO_Europe_2025.docx',
     'title':'SKTT trẻ em/thanh niên WHO châu Âu','authors':'Tarasenko et al. 2025','journal':'Lancet RegEur Q1 IF=15',
     'n':'53 nước','location':'Châu Âu','tool':'Policy review','key':'9 triệu VTN; lo âu>50%; tích hợp trường'},
    {'id':'QT26','region':'Âu-Mỹ-Úc','region_code':'WEST','file_dich':'26_UK_NHS_2025_Parliament.docx','file_tt':'QT26_UK_NHS_2025.docx',
     'title':'Thống kê SKTT Anh — tỷ lệ, dịch vụ, chi tiêu','authors':'Baker 2024','journal':'UK Parliament',
     'n':'quốc gia','location':'England','tool':'NHS Digital','key':'Lo âu 7-16: 3,5%→6,3%; £16 tỷ/năm; TTAD 66,5% cải thiện'},
    {'id':'QT27','region':'Âu-Mỹ-Úc','region_code':'WEST','file_dich':'27_NatureHumanBehav_SocialMedia_2025.docx','file_tt':'QT27_Nature_SocialMedia_2025.docx',
     'title':'MXH ở VTN có/không SKTT — Registered Report','authors':'Fassi et al. 2025','journal':'Nat Hum Behav Q1 IF=30',
     'n':'3340','location':'UK','tool':'DAWBA chẩn đoán','key':'VTN SKTT dùng MXH nhiều hơn g=0,46; nội hóa>ngoại hóa'},
    {'id':'QT32','region':'Âu-Mỹ-Úc','region_code':'WEST','file_dich':'32_Ireland_MyWorld_2024.docx','file_tt':'QT32_Ireland.docx',
     'title':'Xu hướng trầm cảm/lo âu VTN Ireland 2012-2019','authors':'Fitzgerald et al. 2024','journal':'Early Interv Psych',
     'n':'11954','location':'Ireland','tool':'DASS-21','key':'Lo âu tăng 48%; OGA bảo vệ; tự trọng quan trọng hơn'},
    # Bắc Mỹ
    {'id':'QT01','region':'Âu-Mỹ-Úc','region_code':'WEST','file_dich':'01_Jenkins_et_al_2023.docx','file_tt':'',
     'title':'Trầm cảm và lo âu HS THCS đa sắc tộc Mỹ','authors':'Jenkins 2023','journal':'J Early Adolescence',
     'n':'92','location':'California, Mỹ','tool':'PHQ-9A+GAD-10','key':'Lo âu 50,6%; trầm cảm 44%'},
    {'id':'QT23','region':'Âu-Mỹ-Úc','region_code':'WEST','file_dich':'23_JAACAP_US_Trends_2024.docx','file_tt':'QT23_JAACAP_US_2024.docx',
     'title':'Xu hướng RLTT trẻ em Mỹ 2013-2021','authors':'Mojtabai & Olfson 2024','journal':'JAACAP Q1 IF=11',
     'n':'13684154','location':'Mỹ quốc gia','tool':'MH-CLD lâm sàng','key':'Lo âu TĂNG GẤP ĐÔI AOR=2,17; 15-17 tuổi AOR=2,93'},
    {'id':'QT28','region':'Âu-Mỹ-Úc','region_code':'WEST','file_dich':'28_AJP_PediatricAnxiety_2024.docx','file_tt':'QT28_AJP_Treatment_2024.docx',
     'title':'Điều trị lo âu trẻ em — hiện tại và tương lai','authors':'Zugman et al. 2024','journal':'AJP Q1 IF=18',
     'n':'review','location':'NIMH/Mỹ','tool':'Tổng quan','key':'CBT+SSRI 80,7% Walkup 2008 CAMS; NNT=3; ABMT, VR hướng mới'},
    # Úc
    {'id':'QT22','region':'Âu-Mỹ-Úc','region_code':'WEST','file_dich':'22_ScreenTime_2025_BJCP.docx','file_tt':'QT22_ScreenTime_2025.docx',
     'title':'Screen time và trầm cảm/lo âu VTN — dọc 12 tháng','authors':'Li et al. 2025','journal':'BJCP Q1',
     'n':'4058','location':'Úc 134 trường','tool':'PHQ-A+CAS-8','key':'Cắt ngang mạnh, DỌC YẾU; lo âu dọc p=0,443; hai chiều'},
    {'id':'QT25','region':'Âu-Mỹ-Úc','region_code':'WEST','file_dich':'25_EpiPsychSci_2025.docx','file_tt':'QT25_EpiPsychSci_2025.docx',
     'title':'CMH toàn diện VTN Úc 2018-2022','authors':'Crisp et al. 2025','journal':'EpiPsychSci Q1 IF=7',
     'n':'5656','location':'Úc quốc gia','tool':'K10+MHC-SF','key':'Flourishing giảm 53%→44,4%; khoảng cách giới mở rộng'},
    {'id':'QT33','region':'Âu-Mỹ-Úc','region_code':'WEST','file_dich':'33_JAMA_ScreenMedia_2024.docx','file_tt':'QT33_JAMA_Screen.docx',
     'title':'RCT giảm screen time → cải thiện SKTT trẻ em','authors':'Schmidt-Persson 2024','journal':'JAMA NetOpen Q1 IF=13,8',
     'n':'181','location':'Đan Mạch','tool':'SDQ (RCT)','key':'Cohen d=0,53; nội hóa cải thiện mạnh nhất; tuân thủ 97%'},
    # Toàn cầu
    {'id':'QT03','region':'Toàn cầu','region_code':'GLOBAL','file_dich':'04_Zhameden_2025_PLOSONE.docx','file_tt':'',
     'title':'Can thiệp SKTT trường học LMIC','authors':'Zhameden 2025','journal':'PLOS ONE Q1',
     'n':'review','location':'LMIC','tool':'Tổng quan HT','key':'3/4 trầm cảm hiệu quả; 1/4 lo âu; 0 RCT từ VN'},
    {'id':'QT04','region':'Toàn cầu','region_code':'GLOBAL','file_dich':'05_Anderson_2025_Wiley.docx','file_tt':'',
     'title':'Yếu tố gia tăng lo âu VTN','authors':'Anderson 2025','journal':'Wiley',
     'n':'review','location':'Toàn cầu','tool':'Tổng quan','key':'Đa yếu tố: MXH, COVID, giáo dục, gia đình'},
    {'id':'QT29','region':'Toàn cầu','region_code':'GLOBAL','file_dich':'29_CBT_NetworkMeta_2025_BMCPsych.docx','file_tt':'QT29_CBT_NetworkMeta_2025.docx',
     'title':'NMA can thiệp lo âu trẻ em: ACT, CBT, PE, VRET','authors':'Li et al. 2025','journal':'BMC Psychiatry Q1',
     'n':'1711','location':'12 nước','tool':'NMA Bayesian 30 RCT','key':'ACT SUCRA 0,69; CBT 0,66; PE 0,51; VRET 0,51'},
    {'id':'QT30','region':'Toàn cầu','region_code':'GLOBAL','file_dich':'30_GBD_Trends_10-24y_2025.docx','file_tt':'QT30_GBD_Trends.docx',
     'title':'Xu hướng trầm cảm/lo âu 10-24 tuổi 1990-2021 GBD','authors':'Zhang et al. 2025','journal':'J Affect Disord Q1',
     'n':'204 nước','location':'Toàn cầu','tool':'GBD Joinpoint','key':'Lo âu AAPC 0,84%; 10-14 tuổi tăng nhanh nhất 1,44%'},
    {'id':'QT31','region':'Toàn cầu','region_code':'GLOBAL','file_dich':'31_59Countries_Anxiety_2025.docx','file_tt':'QT31_59Countries.docx',
     'title':'Lo âu VTN đi học 59 quốc gia GSHS','authors':'Islam et al. 2025','journal':'J Affect Disord Q1',
     'n':'179937','location':'59 nước LMIC','tool':'GSHS WHO','key':'Thực phẩm AOR=2,22; tự tử AOR=2,84; nữ AOR=1,51; cha mẹ AOR=0,75'},
    {'id':'QT35','region':'Toàn cầu','region_code':'GLOBAL','file_dich':'35_SocialAnxiety_7Countries_2020.docx','file_tt':'QT35_SocialAnxiety_7Countries.docx',
     'title':'Lo âu xã hội ở thanh niên 7 quốc gia (có VN)','authors':'Jefferies & Ungar 2020','journal':'PLOS ONE Q1',
     'n':'6825','location':'7 nước: VN,TQ,ID,RU,TH,US,BR','tool':'SIAS-17','key':'VN SAD=30,7%; tổng 36%; 18% false negatives; không khác giới'},
]

# ===== INSIGHTS from discussions =====
insights = [
    {
        'id': 'INSIGHT_01',
        'region': 'Việt Nam',
        'region_code': 'VN',
        'title': 'Khoảng cách sàng lọc–chẩn đoán tại VN',
        'content': '''PHÁT HIỆN QUAN TRỌNG NHẤT: Tỷ lệ lo âu VTN VN dao động từ 2,3% (chẩn đoán DISC-5, V-NAMHS 2022) đến 86,2% (sàng lọc DASS-21, Bảo Quyên 2025) — khoảng cách gấp 37 lần.
Nguyên nhân CHÍNH: khác biệt CÔNG CỤ ĐO, không phải khác biệt lo âu thực sự.
- DISC-5/DSM-5 (chẩn đoán): 2,3% — con số thấp nhưng chính xác nhất
- GAD-7 (sàng lọc): 40,6% (Hoa 2024 Hà Nội)
- DASS-21 (sàng lọc): 49-86% (tùy mẫu, thời điểm, vùng)
- DASS-Y (sàng lọc VTN riêng): 25,1% (Vĩnh Lộc 2024) — thấp hơn DASS-21 đáng kể
GAP: Cần so sánh GAD-7 vs DASS-21 vs DASS-Y vs DISC-5 trên CÙNG mẫu VTN để xác định ngưỡng phù hợp.'''
    },
    {
        'id': 'INSIGHT_02',
        'region': 'Toàn cầu',
        'region_code': 'GLOBAL',
        'title': 'Xu hướng lo âu VTN tăng toàn cầu — bằng chứng 7 NC dài hạn',
        'content': '''BẰng chứng NHẤT QUÁN từ 7 NC dài hạn: lo âu VTN TĂNG trên toàn cầu.
- Toàn cầu (GBD 204 nước): AAPC 0,84%/năm (1990-2021), tăng tốc từ 2014
- Mỹ (JAACAP): Lo âu TĂNG GẤP ĐÔI AOR=2,17 (2013-2021), 13,7 triệu hồ sơ
- Na Uy: Tăng liên tục 13 năm (2011-2024), n=979.043
- Ireland: Lo âu tăng 48% (2012-2019), TRƯỚC COVID
- Úc: Flourishing giảm 53%→44,4% (2018-2022)
- Anh: Lo âu trẻ 7-16 gấp đôi: 3,5%→6,3%
NGOẠI LỆ DUY NHẤT: Hàn Quốc — cải thiện trước COVID rồi xấu đi sau
VN: THIẾU dữ liệu xu hướng dài hạn — GAP cấp thiết nhất.'''
    },
    {
        'id': 'INSIGHT_03',
        'region': 'Toàn cầu',
        'region_code': 'GLOBAL',
        'title': 'Screen time/MXH và SKTT — tổng hợp 5 NC',
        'content': '''Tác động screen time/MXH lên SKTT VTN TỒN TẠI nhưng PHỨC TẠP hơn giả định.
- Norway 2025 (QT21): MXH giải thích MỘT PHẦN xu hướng tăng (decomposition)
- Nature 2025 (QT27): VTN SKTT dùng MXH nhiều hơn (g=0,46); nội hóa > ngoại hóa
- Li 2025 (QT22): Cắt ngang MẠNH nhưng dọc YẾU (lo âu p=0,443) → quan hệ HAI CHIỀU?
- JAMA 2024 (QT33): RCT giảm ST → cải thiện (d=0,53) → bằng chứng NHÂN QUẢ
- Hoàng Trung Học (VN14): Game điện tử Beta=0,176 tại VN
KẾT LUẬN: Tác động THỰC nhưng CẦN GIẢM ĐỦ MẠNH (>80%) để thấy hiệu quả. VTN nội hóa (lo âu) nhạy cảm hơn ngoại hóa.'''
    },
    {
        'id': 'INSIGHT_04',
        'region': 'Toàn cầu',
        'region_code': 'GLOBAL',
        'title': 'Can thiệp hiệu quả cho lo âu VTN — bằng chứng',
        'content': '''Tổng hợp can thiệp từ 5 NC:
1. CBT+SSRI: 80,7% đáp ứng (Walkup 2008 NEJM — CAMS), NNT=3 — HIỆU QUẢ NHẤT
2. CBT đơn: 59,7% đáp ứng, effect size 0,31 (CAMS)
3. ACT: SUCRA 0,69 — xếp hạng 1 NMA nhưng ít RCT (Li et al. 2025 BMC)
4. CBT nhóm: SUCRA 0,66 — bằng chứng mạnh nhất (16/30 RCT)
5. Hoạt động thể chất (PE): SUCRA 0,51 — DỄ triển khai tại trường
6. Giảm screen time: Cohen d=0,53 chỉ 2 tuần (JAMA 2024 RCT)
7. Hỗ trợ trường: OR=0,562 bảo vệ (Wen 2020)
VN: CHƯA CÓ RCT nào — GAP lớn nhất. CBT nhóm + PE tại trường khả thi nhất.'''
    },
    {
        'id': 'INSIGHT_05',
        'region': 'Việt Nam',
        'region_code': 'VN',
        'title': 'Yếu tố nguy cơ và bảo vệ ở VTN VN — tổng hợp',
        'content': '''YẾU TỐ NGUY CƠ (VN + liên bài):
- Áp lực học tập: OR=11,58 (Wen 2020 TQ); ESSA≥59 tăng nguy cơ (Vĩnh Lộc 2024 VN)
- Giấc ngủ <5h: AOR=13,71 (Zhu 2025 TQ) — yếu tố mạnh nhất
- Bất an thực phẩm: AOR=2,22 (Islam 2025, 59 nước LMIC)
- DTTS/vùng khó: 54,4% DTTS Lạng Sơn (Ngô Anh Vinh 2024)
- ACEs (nghịch cảnh thời thơ ấu): Coef=0,28 (Ngô Anh Vinh 2024)
- Cha mẹ ly hôn: tăng nguy cơ (Vĩnh Lộc 2024; Zhu 2025)
- Nữ giới: AOR=1,51 (59 Countries); nhất quán trừ lo âu XH
YẾU TỐ BẢO VỆ:
- Hỗ trợ trường: OR=0,562 (Wen 2020)
- Cha mẹ tham gia: AOR=0,75 (Islam 2025)
- Lạc quan: β gián tiếp=−0,24 qua lo âu (Thảo Vi 2025 Huế)
- One Good Adult: nhất quán 2 đợt Ireland (Fitzgerald 2024)'''
    },
    {
        'id': 'INSIGHT_06',
        'region': 'Việt Nam',
        'region_code': 'VN',
        'title': 'Lo âu xã hội ở VN — phát hiện đặc biệt',
        'content': '''Jefferies & Ungar 2020 (QT35): VN SAD = 30,7% — cứ 3 thanh niên VN có 1 có SAD.
ĐẶC BIỆT:
- KHÔNG khác biệt giới tính (trái với lo âu tổng quát nữ>nam)
- 18% "false negatives" — có SAD nhưng KHÔNG nhận ra → "lo âu ẩn"
- VN: nói với người có thẩm quyền xếp hạng lo âu 3/17 — văn hóa thứ bậc
- SIAS-17 đo lo âu TƯƠNG TÁC (không phải biểu diễn) → có thể bỏ sót lo âu thi cử
- So sánh ĐNA: VN 30,7% vs Indonesia 22,9% vs Thái Lan 41,4% — cùng ĐNA nhưng rất khác
GAP: VN chưa có NC lo âu XH chuyên biệt ở THCS/THPT (bài này 16-29 tuổi).'''
    },
    {
        'id': 'INSIGHT_07',
        'region': 'Đông Bắc Á',
        'region_code': 'NEA',
        'title': 'Hàn Quốc — mô hình đảo chiều duy nhất',
        'content': '''Cho et al. 2024 (QT34): KYRBS 16 năm (2006-2022) — dài nhất trong Đề tài.
MÔ HÌNH ĐẢO CHIỀU (duy nhất):
- Trước COVID (2006-2019): SKTT CẢI THIỆN liên tục (β = −1,41)
- Sau COVID (2020-2022): XẤU ĐI đột ngột (β = +2,39)
- Thay đổi: βdiff = 3,80***
BẤT BÌNH ĐẲNG THU NHẬP MỞ RỘNG sau COVID:
- Stress nghèo: 62,77% vs giàu: 40,07% — chênh 22,7 điểm
- COVID tác động KHÔNG ĐỀU — nhóm nghèo bị nặng hơn
SO SÁNH: Phương Tây (Na Uy, Mỹ, Ireland) tăng LIÊN TỤC, không đảo chiều.
HÀM Ý VN: (1) Can thiệp CÓ THỂ hiệu quả (giai đoạn cải thiện HQ), (2) VN cần phân tầng SKTT theo thu nhập.'''
    },
    {
        'id': 'INSIGHT_08',
        'region': 'Toàn cầu',
        'region_code': 'GLOBAL',
        'title': 'Giới tính — nhất quán xuyên suốt (trừ lo âu XH)',
        'content': '''NỮ > NAM trong lo âu tổng quát — NHẤT QUÁN:
- Hoa 2024 VN: GAD-7 M nữ 1,74 vs nam 1,50 (p<0,01)
- JAACAP 2024 Mỹ: Nữ 22,3% vs nam 16,5%
- 59 Countries: AOR=1,51 (p<0,001) — 59 nước
- Úc CMH: Flourishing nữ 36,3% vs nam 52,2% — chênh 15,9 điểm
- Ireland: Nữ tăng nhanh hơn nam — khoảng cách MỞ RỘNG
NGOẠI LỆ: Lo âu XÃ HỘI (Jefferies 2020 QT35): Nam 35,6% = Nữ 36,5% (n.s.)
→ Lo âu XH ảnh hưởng nam = nữ — cơ chế KHÁC lo âu tổng quát.'''
    },
    {
        'id': 'INSIGHT_09',
        'region': 'Việt Nam',
        'region_code': 'VN',
        'title': 'Đề cương nghiên cứu đề xuất — 3 giai đoạn',
        'content': '''ĐỀ CƯƠNG GỢI Ý dựa trên 35 bài NC:
GIAI ĐOẠN 1 — Khảo sát (6 tháng):
- GAD-7 + DASS-Y + SIAS-17, 1.000+ HS THCS/THPT, 3 vùng (đô thị/nông thôn/DTTS)
- Thêm: ESSA áp lực, ACEs, screen time, giấc ngủ
- Mẫu con n≈300: DISC-5 chẩn đoán → xác định khoảng cách sàng lọc–chẩn đoán
GIAI ĐOẠN 2 — Can thiệp RCT (12 tháng):
- CBT nhóm + PE tại trường (BMC NMA: cả hai hiệu quả)
- Giảm screen time toàn gia đình (JAMA 2024: d=0,53)
- 5 trường can thiệp vs 5 đối chứng (cluster RCT)
GIAI ĐOẠN 3 — Phỏng vấn sâu (6 tháng):
- "Lo âu ẩn" (18% false negatives — Jefferies 2020)
- Rào cản tiếp cận dịch vụ (V-NAMHS: chỉ 8,4%)'''
    },
    {
        'id': 'INSIGHT_10',
        'region': 'Việt Nam',
        'region_code': 'VN',
        'title': 'Top 10 khoảng trống nghiên cứu — ưu tiên cho VN',
        'content': '''1. RCT can thiệp SKTT tại trường VN (0 RCT hiện có) — RẤT CAO
2. So sánh sàng lọc vs chẩn đoán trên cùng mẫu (chênh 17-37 lần) — RẤT CAO
3. Khảo sát quốc gia lặp lại (xu hướng dài hạn — VN chưa có) — RẤT CAO
4. Phân tích SKTT VTN theo thu nhập (Korea: chênh 22,7 điểm) — CAO
5. NC lo âu XÃ HỘI chuyên biệt ở VTN VN (SAD=30,7% nhưng 18% không biết) — CAO
6. Decomposition nguyên nhân xu hướng tại VN (như Norway) — CAO
7. RCT giảm screen time tại trường VN (0 RCT) — CAO
8. Đánh giá DASS-Y vs DASS-21 vs GAD-7 ở VTN VN (ngưỡng nào phù hợp?) — TB
9. NC SKTT VTN DTTS/vùng khó (54,4% DTTS vs 40,6% Hà Nội) — CAO
10. Bất an thực phẩm → lo âu VTN VN (AOR=2,22 toàn cầu) — TB'''
    },
    {
        'id': 'CROSS_STUDY',
        'region': 'Toàn cầu',
        'region_code': 'GLOBAL',
        'title': 'Tổng hợp liên bài — Cross-Study Synthesis 35 NC',
        'content': read_docx(os.path.join(BASE, '..', 'Tổng hợp liên bài báo - Lo âu HS - 04042026.docx'), max_chars=5000),
    },
]

# ===== BUILD RAG =====
print('Rebuilding RAG...')
client = chromadb.PersistentClient(path=RAG_DIR)

# Delete and recreate
try:
    client.delete_collection('lo_au_full')
except:
    pass

col = client.create_collection(
    name='lo_au_full',
    metadata={'description': 'Lo âu VTN — 35 bài NC + 11 insights, phân loại theo vùng'}
)

documents = []
metadatas = []
ids = []

# Add papers
for p in papers:
    # Read translation + summary
    dich_text = ''
    if p['file_dich']:
        dich_path = os.path.join(DICH_DIR, p['file_dich'])
        if os.path.exists(dich_path):
            dich_text = read_docx(dich_path, max_chars=3000)

    tt_text = ''
    if p['file_tt']:
        tt_path = os.path.join(TT_DIR, p['file_tt'])
        if os.path.exists(tt_path):
            tt_text = read_docx(tt_path, max_chars=2000)

    combined = f"""BÀI {p['id']} [{p['region']}]: {p['title']}
Tác giả: {p['authors']} | Tạp chí: {p['journal']}
n={p['n']} | Địa bàn: {p['location']} | Công cụ: {p['tool']}
Vùng: {p['region']} ({p['region_code']})
Phát hiện chính: {p['key']}

--- TÓM TẮT ---
{tt_text[:1500]}

--- BẢN DỊCH (trích) ---
{dich_text[:2000]}"""

    documents.append(combined)
    metadatas.append({
        'paper_id': p['id'],
        'title': p['title'],
        'authors': p['authors'],
        'journal': p['journal'],
        'n': str(p['n']),
        'location': p['location'],
        'tool': p['tool'],
        'key_finding': p['key'],
        'region': p['region'],
        'region_code': p['region_code'],
        'type': 'paper',
    })
    ids.append(p['id'])

# Add insights
for ins in insights:
    documents.append(f"INSIGHT [{ins['region']}]: {ins['title']}\n\n{ins['content']}")
    metadatas.append({
        'paper_id': ins['id'],
        'title': ins['title'],
        'authors': 'Tổng hợp từ trao đổi',
        'journal': '',
        'n': '',
        'location': ins['region'],
        'tool': '',
        'key_finding': ins['title'],
        'region': ins['region'],
        'region_code': ins['region_code'],
        'type': 'insight',
    })
    ids.append(ins['id'])

col.add(documents=documents, metadatas=metadatas, ids=ids)
print(f'Added {len(documents)} entries ({len(papers)} papers + {len(insights)} insights)')
print(f'Collection count: {col.count()}')

# ===== VERIFY =====
print('\n=== VERIFICATION ===')

# Test queries
tests = [
    ('lo âu VTN Việt Nam tỷ lệ', 'VN'),
    ('can thiệp CBT trường học', 'GLOBAL'),
    ('screen time mạng xã hội tác động', 'GLOBAL'),
    ('Hàn Quốc xu hướng COVID thu nhập', 'NEA'),
    ('lo âu xã hội SAD thanh niên', 'GLOBAL'),
    ('DASS-21 vs GAD-7 sàng lọc chẩn đoán', 'VN'),
    ('giấc ngủ yếu tố nguy cơ', 'NEA'),
    ('Ireland One Good Adult OGA', 'WEST'),
]

for query, expected_region in tests:
    results = col.query(query_texts=[query], n_results=3)
    top = results['metadatas'][0][0]
    top_region = top.get('region_code', '?')
    match = '✓' if top_region == expected_region else '~'
    print(f'  {match} "{query[:40]}..." → [{top["paper_id"]}] {top["title"][:40]}... ({top["region"]})')

# Count by region
print('\n=== BY REGION ===')
all_meta = col.get(include=['metadatas'])
region_counts = {}
for m in all_meta['metadatas']:
    r = m.get('region', '?')
    t = m.get('type', 'paper')
    key = f'{r} ({t})'
    region_counts[key] = region_counts.get(key, 0) + 1
for k, v in sorted(region_counts.items()):
    print(f'  {k}: {v}')

print(f'\nTOTAL: {col.count()} entries')
print('DONE!')
