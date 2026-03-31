# -*- coding: utf-8 -*-
"""
Build/Rebuild RAG index từ tất cả papers + translations + manual entries
Chạy: python build_rag.py
Sao lưu: copy rag_db/ folder
"""
import sys, io, os, glob, re, json, shutil
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass

import chromadb
from sentence_transformers import SentenceTransformer

# Config
MODEL_NAME = os.environ.get('EMBEDDING_MODEL', 'all-MiniLM-L6-v2')
RAG_DB = 'rag_db'
COLLECTION = 'lo_au_v3'
PAPERS_DIR = os.path.join('..', 'papers')
TRANSLATIONS_DIR = os.path.join('..', 'Translations')
BAN_DICH_DIR = os.path.join('..', '03_Ban-dich')
DOCS_DIR = os.path.join('web', 'docs')

print(f'Model: {MODEL_NAME}')
model = SentenceTransformer(MODEL_NAME)

# Backup old DB
if os.path.exists(RAG_DB):
    backup = f'{RAG_DB}_backup'
    if os.path.exists(backup):
        shutil.rmtree(backup)
    shutil.copytree(RAG_DB, backup)
    print(f'Backup: {backup}/')

client = chromadb.PersistentClient(path=RAG_DB)
try: client.delete_collection(COLLECTION)
except: pass
col = client.create_collection(COLLECTION, metadata={"hnsw:space": "cosine"})

def chunk_text(text, chunk_size=400, overlap=80):
    words = text.split()
    chunks = []
    page = '1'
    for i in range(0, len(words), chunk_size - overlap):
        ch = ' '.join(words[i:i+chunk_size])
        if len(ch.strip()) < 50: continue
        pm = re.findall(r'Trang\s+(\d+[\u2013-]?\d*)', ch)
        if pm: page = pm[-1]
        pm2 = re.findall(r'---\s*Trang\s+([\w\d/\u2013-]+)', ch)
        if pm2: page = pm2[-1]
        chunks.append((ch, page))
    return chunks

all_chunks = []
all_ids = []
all_meta = []
doc_id = 0

# 1. PDFs
for pdf_dir in [PAPERS_DIR, os.path.join('..', '02_Papers-goc'),
                os.path.join('..', '02_Papers-goc', 'Viet-Nam'),
                os.path.join('..', '02_Papers-goc', 'Dong-Nam-A'),
                os.path.join('..', '02_Papers-goc', 'The-gioi'),
                os.path.join('..', 'Originals')]:
    if not os.path.exists(pdf_dir): continue
    for pdf in sorted(glob.glob(os.path.join(pdf_dir, '*.pdf'))):
        fname = os.path.basename(pdf)
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(pdf)
            text = ''
            for p in reader.pages[:50]:
                text += (p.extract_text() or '') + ' '
            if len(text.strip()) < 200: continue
            for ci, (ch, pg) in enumerate(chunk_text(text)):
                all_chunks.append(ch)
                all_ids.append(f'd{doc_id}_c{ci}')
                all_meta.append({'source': fname, 'type': 'pdf', 'chunk_index': ci, 'page_ref': pg})
            doc_id += 1
        except: pass

# 2. Translations (MD)
for md_dir in [TRANSLATIONS_DIR, BAN_DICH_DIR, PAPERS_DIR, DOCS_DIR, '..']:
    if not os.path.exists(md_dir): continue
    for md in sorted(glob.glob(os.path.join(md_dir, '*.md'))):
        fname = os.path.basename(md)
        if 'MEMORY' in fname or 'PROMPT' in fname or 'HUONG_DAN' in fname: continue
        try:
            with open(md, 'r', encoding='utf-8') as f:
                text = f.read()
            if len(text.strip()) < 200: continue
            for ci, (ch, pg) in enumerate(chunk_text(text)):
                all_chunks.append(ch)
                all_ids.append(f'd{doc_id}_c{ci}')
                all_meta.append({'source': fname, 'type': 'markdown', 'chunk_index': ci, 'page_ref': pg})
            doc_id += 1
        except: pass

# 3. DOCX translations (03_Ban-dich)
if os.path.exists(BAN_DICH_DIR):
    for docx in sorted(glob.glob(os.path.join(BAN_DICH_DIR, '0*.docx'))):
        fname = os.path.basename(docx)
        try:
            from docx import Document
            doc = Document(docx)
            text = ' '.join(p.text for p in doc.paragraphs)
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        text += ' ' + cell.text
            if len(text.strip()) < 200: continue
            for ci, (ch, pg) in enumerate(chunk_text(text)):
                all_chunks.append(ch)
                all_ids.append(f'd{doc_id}_c{ci}')
                all_meta.append({'source': fname, 'type': 'docx', 'chunk_index': ci, 'page_ref': pg})
            doc_id += 1
        except: pass

print(f'Chunks from files: {len(all_chunks)} from {doc_id} docs')

# 4. Manual entries (verified data, NOT fabricated)
manual = [
    # 11 bài tuần trước
    'Jenkins và cs. (2023): 75 HS THCS đa sắc tộc tại San Diego, Mỹ. PHQ-9A + GAD-10. Lo âu 50,6%, trầm cảm 44%. Nữ > nam (p=0,002). Bạo lực giới, COVID-19, phân biệt chủng tộc.',
    'Saikia và cs. (2023): 360 HS Assam, Ấn Độ. DASS-21. Lo âu 24,4%. Nam > nữ (30% vs 18,9%, P=0,049) — trái y văn. Cha mẹ đơn thân, rượu, game.',
    'Mandaknalli & Malusare (2021): 450 HS Ấn Độ. Lo âu nhẹ 49,4%, TB 43,3%, nặng 7,3%. Nữ nặng hơn (10,9% vs 3,8%). Giấc ngủ, vận động.',
    'NSCH 2020: Hoa Kỳ quốc gia, n=55.162. Lo âu 16,1% (chẩn đoán). Tăng 61% (2016-2023). Nữ 20,1% vs nam 12,3%. Khó kết bạn gấp 10 lần.',
    'Alharbi và cs. (2019): 1.245 HS THPT Ả Rập Saudi. PHQ-9 + GAD-7. Lo âu 63,5%, trầm cảm 74% (ngưỡng thấp). Nữ > nam (P<0,001).',
    'Nakie và cs. (2022): 849 HS Ethiopia. DASS-21. Lo âu 66,7%, trầm cảm 41,4%, stress 52,2%. Khat AOR=5,6, thuốc AOR=4,8. BMC Psychiatry Q1.',
    'Chen và cs. (2023): 63.205 HS miền Tây TQ. PHQ-9 + GAD-7. Lo âu 13,9%, trầm cảm 23%. Bắt nạt, giấc ngủ, game. BMC Psychiatry Q1.',
    'Wen và cs. (2020): 900 HS THCS nông thôn TQ. LPA 3 nhóm lo âu. Lo âu nặng 24,78%. Nam > nữ. Hỗ trợ trường bảo vệ.',
    'Qiu và cs. (2022): 2.079 HS THCS TQ. Lo âu 13,4%, trầm cảm 26%. Nuôi dạy tiêu cực OR=2,01. Phục hồi thấp OR=6,74. Frontiers Q2.',
    'Xu và cs. (2021): 373.216 HS TQ — LỚN NHẤT TOÀN CẦU. GAD-7. Lo âu 9,89%. Nam > nữ (10,11% vs 9,66%). Nông thôn 12,80%. J Affect Disord Q1.',
    'Bhardwaj và cs. (2020): 288 HS trường công Chandigarh, Ấn Độ. DASS-21. Lo âu 73,3% — CAO NHẤT. Trầm cảm 64,9%, stress 74,7%.',
    # 10 bài tuần này
    'Hoa và cs. (2024): 3.910 HS THPT Hà Nội. GAD-7 (α=0,916). Lo âu 40,6%. Nữ 1,74 vs nam 1,50 (p<0,01). Phỏng vấn sâu: áp lực thi, kỳ thị, gia đình.',
    'V-NAMHS (2022): 5.996 VTN 10-17 tuổi, 38 tỉnh VN. DISC-5/DSM-5. Vấn đề SKTT 21,7%, lo âu 18,6%, rối loạn lo âu 2,3%. Chỉ 8,4% tiếp cận dịch vụ.',
    'Ngô Anh Vinh (2024): 845 HS DTTS nội trú Lạng Sơn. DASS-21. Lo âu 54,4%, trầm cảm 59%. ACEs 48,9%. Tình bạn kém tăng lo âu.',
    'GBD ASEAN (2025): 10 nước ASEAN, Lancet. 80,4 triệu ca RLTT, tăng 70% (1990-2021). VN 10,1%, Malaysia 13,2%. 10-14 tuổi: 16,3% DALYs.',
    'Zhameden (2025): 6 RCTs tại LMIC, 1.587 HS. CBT 3/4 trầm cảm nhưng 1/4 lo âu. GRADE rất thấp. Không có RCT từ VN. PLOS ONE.',
    'Anderson (2025): Tổng quan 61 bài. 31,9% VTN có lo âu. Gen Z cao nhất. 48/52 NC: áp lực học tập gây SKTT kém. Wiley.',
    'Zhu (2025): 9.831 HS THCS/THPT Suzhou, TQ. PHQ-9. Trầm cảm 14,5%+5,8%. Ngủ <5h: AOR=13,710. Ngoài trời AOR=0,557. BMC.',
    'Puyat (2025): 30.127 TN Philippines. CES-D-11. Trầm cảm 9,6% (2013) → 20,9% (2021). LGBTQ+ 32,3%. Nữ 24,3%. Nghèo 25,1%.',
    'Mudunna (2025): Nam Á 8 nước, tổng quan hệ thống. Tỷ lệ 1,5-81,6%. Kỳ thị SKTT rào cản lớn nhất. Lancet Regional Health SEA.',
    'Pham (2024): 273+273 VTN Huế. Chăm sóc cảm xúc β=-0,40 cho lo âu. Thể chất không ý nghĩa. Cảm xúc > thể chất.',
    # Cross-study synthesis
    'So sánh sàng lọc vs chẩn đoán lo âu VN: GAD-7 40,6% (Hoa), DASS-21 54,4% (Ngô), DISC-5 chỉ 2,3% (V-NAMHS). Chênh 17-24 lần.',
    'Giới tính: 8/18 bài nữ > nam (y văn). 3 ngoại lệ: Saikia, Wen, Xu (nam > nữ). V-NAMHS không khác biệt.',
    'Giấc ngủ: AOR=13,71 cho ngủ <5h (Zhu TQ). Hoa VN: lo lắng ảnh hưởng giấc ngủ. Hoàng TH: giấc ngủ kém là yếu tố nguy cơ.',
    'Gia đình: chăm sóc cảm xúc β=-0,40 (Pham VN). Có người tâm sự OR=0,22 (Dong TQ). Đơn thân AOR=1,43, tái hôn 1,84 (Zhu TQ).',
    'Can thiệp: 0 RCT từ VN. CBT kém cho lo âu (1/4). GRADE rất thấp. Cần RCT trường học VN.',
    'Dịch vụ: VN 8,4% tiếp cận. Philippines 2%. Phụ huynh VN 5,1% nhận ra. Bác sĩ ASEAN 0,2-4,3/100K.',
    '10 gap ưu tiên VN: (1) RCT trường học (2) GAD-7 vs DISC-5 cùng mẫu (3) NC dọc COVID (4) Giấc ngủ-lo âu (5) Đào tạo phụ huynh.',
    'GAD-7: thang sàng lọc lo âu 7 mục, điểm 0-21. Nhẹ 5-9, TB 10-14, nặng 15-21. Spitzer 2006.',
    'DASS-21: 21 mục, 3 thang con (trầm cảm, lo âu, stress). Lovibond 1995.',
    'DISC-5: chẩn đoán theo DSM-5, phỏng vấn cấu trúc. Tỷ lệ thấp hơn sàng lọc 10-37 lần.',
    'PHQ-9: sàng lọc trầm cảm 9 mục, điểm 0-27. Spitzer 1999.',
]

for i, entry in enumerate(manual):
    all_chunks.append(entry)
    all_ids.append(f'manual_{i}')
    all_meta.append({'source': 'manual_entry', 'type': 'manual', 'page_ref': 'tổng hợp', 'chunk_index': i})

print(f'Total chunks: {len(all_chunks)}')

# Embed and store in batches
batch = 100
for i in range(0, len(all_chunks), batch):
    bc = all_chunks[i:i+batch]
    bi = all_ids[i:i+batch]
    bm = all_meta[i:i+batch]
    embs = model.encode(bc).tolist()
    col.add(documents=bc, embeddings=embs, ids=bi, metadatas=bm)
    print(f'  Batch {i//batch+1}/{(len(all_chunks)-1)//batch+1}')

print(f'DONE: {col.count()} chunks indexed with model {MODEL_NAME}')
print(f'DB size: {os.path.getsize("rag_db/chroma.sqlite3") // 1024 // 1024}MB')
